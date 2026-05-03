import asyncio
import logging
import os
import time
import base64
import sqlite3
import subprocess
import json
import glob
from io import BytesIO
from datetime import datetime, timedelta
from PIL import Image

from aiogram import Bot, Dispatcher, Router, types, F
from aiogram.filters import Command
from aiogram.types import FSInputFile, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from dotenv import load_dotenv
from google import genai
from google.genai import types as genai_types
from aiohttp import web

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = Router()

ADMIN_USER_ID = int(os.getenv("ADMIN_USER_ID", "7361263893"))
DEVELOPER_NAME = "Omar Abd El Gawaad"
DEVELOPER_USERNAME = "@omarawad68"

user_conversion_choice = {}
user_pending_file = {}

SYSTEM_PROMPT = """
أنت "مستشار الذكاء الاصطناعي الخارق". أنت تجمع بين خبير موسوعي ومبرمج عبقري. هدفك تقديم إجابات دقيقة واحترافية في كل المجالات، مع قدرة استثنائية على البرمجة.

قواعدك الصارمة:
0.  **اللغة التلقائية:** **يجب عليك الرد بنفس لغة سؤال المستخدم.**
1.  **ممنوع المقدمات:** لا تبدأ أي إجابة بعبارات مثل "أهلاً بك"، "بصفتي...". ابدأ الإجابة مباشرة.
2.  **تنسيق الردود:** للقوائم استخدم "- ". للكود استخدم ``` مع تحديد اللغة. لا تستخدم "#" أبداً.
3.  **الهوية المزدوجة:** إذا كان السؤال برمجياً ركز على الكود، وإذا كان عاماً قدم شرحاً مباشراً.
4.  **تحليل المستندات والصور:** عند تحليل مستند أو صورة، ابدأ مباشرة بتحليل المحتوى بدون مقدمات.
5.  **الأمان:** ترفض أي طلب لإنشاء محتوى ضار أو غير قانوني.
"""

def init_db():
    try:
        conn = sqlite3.connect('bot_stats.db')
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY, username TEXT, first_name TEXT,
            last_name TEXT, joined_date TEXT, last_active TEXT, total_messages INTEGER DEFAULT 0)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS daily_stats (
            date TEXT PRIMARY KEY, active_users INTEGER DEFAULT 0, total_messages INTEGER DEFAULT 0)''')
        conn.commit()
        conn.close()
        logger.info("Database initialized")
    except Exception as e:
        logger.error(f"DB init error: {e}")

def update_user_activity(user: types.User):
    try:
        conn = sqlite3.connect('bot_stats.db')
        cursor = conn.cursor()
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        today = datetime.now().strftime("%Y-%m-%d")
        cursor.execute('''INSERT INTO users (user_id, username, first_name, last_name, joined_date, last_active, total_messages)
            VALUES (?, ?, ?, ?, ?, ?, 1) ON CONFLICT(user_id) DO UPDATE SET
            username=excluded.username, first_name=excluded.first_name, last_name=excluded.last_name,
            last_active=excluded.last_active, total_messages=users.total_messages+1''',
            (user.id, user.username, user.first_name, user.last_name, now, now))
        cursor.execute('''INSERT INTO daily_stats (date, active_users, total_messages) VALUES (?, 1, 1)
            ON CONFLICT(date) DO UPDATE SET active_users=active_users+1, total_messages=total_messages+1''', (today,))
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"User activity error: {e}")

class AsyncGeminiClient:
    def __init__(self, model: str = "gemini-3.1-flash-lite-preview"):
PyMuPDFFself.client = genai.Client()
        self.model = model
        self.conversations = {}

    async def generate(self, prompt: str, user_id: str = "default") -> str:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._sync_generate, prompt, user_id)

    def _sync_generate(self, prompt: str, user_id: str = "default") -> str:
        if user_id not in self.conversations:
            self.conversations[user_id] = []
        
        self.conversations[user_id].append({
            "role": "user",
            "parts": [{"text": prompt}]
        })
        
        if len(self.conversations[user_id]) > 15:
            self.conversations[user_id] = self.conversations[user_id][-15:]
        
        full_context = [
            {"role": "user", "parts": [{"text": "أنت مستشار ذكي. تذكر محادثتنا."}]},
            {"role": "model", "parts": [{"text": "حسناً، سأتذكر محادثتنا."}]},
            *self.conversations[user_id]
        ]
        
        for attempt in range(3):
            try:
                response = self.client.models.generate_content(
                    model=self.model,
                    contents=full_context,
                    config=genai_types.GenerateContentConfig(system_instruction=SYSTEM_PROMPT))
                
                reply = response.text
                
                self.conversations[user_id].append({
                    "role": "model",
                    "parts": [{"text": reply}]
                })
                
                return reply
            except Exception as e:
                logger.error(f"Gemini error (attempt {attempt+1}): {e}")
                if attempt < 2:
                    time.sleep(1)
                else:
                    return "عذراً، حدث خطأ مؤقت."

    async def generate_with_media(self, prompt: str, media_parts: list) -> str:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._sync_generate_with_media, prompt, media_parts)

    def _sync_generate_with_media(self, prompt: str, media_parts: list) -> str:
        for attempt in range(3):
            try:
                contents = media_parts + [{"text": prompt}]
                response = self.client.models.generate_content(
                    model=self.model, contents=contents,
                    config=genai_types.GenerateContentConfig(system_instruction=SYSTEM_PROMPT))
                return response.text
            except Exception as e:
                logger.error(f"Gemini media error (attempt {attempt+1}): {e}")
                if attempt < 2: time.sleep(1)
                else: return "عذراً، حدث خطأ مؤقت."

gemini_client = AsyncGeminiClient()

def convert_image_to_png(image_bytes: bytes):
    try:
        img = Image.open(BytesIO(image_bytes))
        fmt = img.format
        if fmt not in ['JPEG','PNG','GIF']:
            buf = BytesIO()
            img.convert('RGB').save(buf, format='PNG')
            return buf.getvalue(), "image/png"
        mime = f"image/{fmt.lower()}"
        return image_bytes, "image/jpeg" if mime=="image/jpg" else mime
    except:
        return image_bytes, "image/jpeg"

def create_docx_file(text: str, filepath: str):
    from docx import Document
    doc = Document()
    doc.add_heading('مستند تم إنشاؤه بواسطة البوت', level=1)
    doc.add_paragraph(text)
    doc.save(filepath)

def create_pdf_file(text: str, filepath: str):
    docx_path = filepath.replace('.pdf', '.docx')
    create_docx_file(text, docx_path)
    run_libreoffice(['--convert-to', 'pdf', '--outdir', os.path.dirname(filepath), docx_path])
    os.remove(docx_path)

def create_excel_file(text: str, filepath: str):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    
    wb = Workbook()
    ws = wb.active
    ws.title = "البيانات"
    
    lines = text.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]
    
    if len(lines) >= 2:
        headers = [h.strip() for h in lines[0].replace('،', ',').split(',')]
        data_rows = []
        for line in lines[1:]:
            row = [cell.strip() for cell in line.replace('،', ',').split(',')]
            data_rows.append(row)
        
        header_font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        cell_font = Font(name='Arial', size=12)
        cell_alignment = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        if headers:
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        title_cell = ws['A1']
        title_cell.value = "مستند تم إنشاؤه بواسطة البوت"
        title_cell.font = Font(name='Arial', size=16, bold=True, color='2F5496')
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        for row_idx, row_data in enumerate(data_rows, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.font = cell_font
                cell.alignment = cell_alignment
                cell.border = thin_border
                if row_idx % 2 == 0:
                    cell.fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
        
        for col in ws.columns:
            max_length = 0
            column_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 4, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        if headers:
            ws.auto_filter.ref = ws.dimensions
    
    else:
        ws['A1'] = "النص المحول"
        ws['A1'].font = Font(name='Arial', size=14, bold=True, color='2F5496')
        for row_idx, line in enumerate(lines, 2):
            ws.cell(row=row_idx, column=1, value=line)
        ws.column_dimensions['A'].width = 50
    
    wb.save(filepath)

def detect_conversion_intent(text: str):
    text_lower = text.lower()
    
    excel_patterns = [
        "حولي النص التالي لملف اكسيل", "حولي لملف اكسيل", "حول لملف اكسيل",
        "ملف اكسيل", "ملف excel", "اكسيل", "excel", "xlsx",
        "اعملي ملف اكسيل", "اعمل ملف excel",
    ]
    
    word_patterns = [
        "حولي النص التالي لملف وورد", "حولي لملف وورد", "حول لملف وورد",
        "ملف وورد", "ملف word", "وورد", "word", "docx",
        "اعملي ملف وورد", "اعمل ملف word",
    ]
    
    pdf_patterns = [
        "حولي النص التالي لملف pdf", "حولي لملف pdf", "حول لملف pdf",
        "ملف pdf", "بي دي اف", "pdf",
        "اعملي ملف pdf", "اعمل ملف بي دي اف",
    ]
    
    for pattern in excel_patterns:
        if pattern in text_lower:
            idx = text_lower.find(pattern)
            content = text[idx + len(pattern):].strip()
            if not content:
                content = text[:idx].strip()
                for prefix in ["حولي", "حول", "حوّل", "خلي", "اعمل", "سوي", "ابعتلي", "انزلي", "حملي"]:
                    if content.startswith(prefix):
                        content = content[len(prefix):].strip()
                        break
            if content: return "excel", content
            else: return "EXCEL_NEED_TEXT", ""
    
    for pattern in word_patterns:
        if pattern in text_lower:
            idx = text_lower.find(pattern)
            content = text[idx + len(pattern):].strip()
            if not content:
                content = text[:idx].strip()
                for prefix in ["حولي", "حول", "حوّل", "خلي", "اعمل", "سوي", "ابعتلي", "انزلي", "حملي"]:
                    if content.startswith(prefix):
                        content = content[len(prefix):].strip()
                        break
            if content: return "docx", content
            else: return "WORD_NEED_TEXT", ""
    
    for pattern in pdf_patterns:
        if pattern in text_lower:
            idx = text_lower.find(pattern)
            content = text[idx + len(pattern):].strip()
            if not content:
                content = text[:idx].strip()
                for prefix in ["حولي", "حول", "حوّل", "خلي", "اعمل", "سوي", "ابعتلي", "انزلي", "حملي"]:
                    if content.startswith(prefix):
                        content = content[len(prefix):].strip()
                        break
            if content: return "pdf", content
            else: return "PDF_NEED_TEXT", ""
    
    return None, None

def get_conversion_keyboard():
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 Word → PDF", callback_data="convert_word2pdf"),
         InlineKeyboardButton(text="📄 PDF → Word", callback_data="convert_pdf2word")],
        [InlineKeyboardButton(text="📊 Excel → PDF", callback_data="convert_excel2pdf"),
         InlineKeyboardButton(text="📊 PDF → Excel", callback_data="convert_pdf2excel")],
        [InlineKeyboardButton(text="📊 Excel → Word", callback_data="convert_excel2word"),
         InlineKeyboardButton(text="📄 Word → Excel", callback_data="convert_word2excel")],
        [InlineKeyboardButton(text="🔄 أي صيغة لأي صيغة", callback_data="convert_any")]
    ])
    return keyboard


def run_libreoffice(args, timeout=60):
    """
    تحويل PDF إلى Word باستخدام LibreOffice مع فلتر writer_pdf_import.
    يحافظ هذا الفلتر على الجداول والتنسيق الأصلي.
    """
    full_args = ['libreoffice', '--headless', '-env:UserInstallation=file:///tmp/libreoffice']
    input_files = [a for a in args if os.path.exists(a) and a.lower().endswith('.pdf')]
    if input_files:
        full_args.append('--infilter="writer_pdf_import"')
    
    full_args.extend(args)
    return subprocess.run(
        full_args,
        capture_output=True, text=True, timeout=timeout,
        env={**os.environ, 'HOME': '/tmp', 'USERPROFILE': '/tmp'}
    )


def convert_pdf_to_excel(input_path: str, output_path: str):
    """
    تحويل PDF إلى Excel مع ضبط الورقة بأكملها من اليمين لليسار.
    """
    import pdfplumber
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Data"
    ws.sheet_view.rightToLeft = True  # <--- هذا هو الحل الجذري لاتجاه النص العربي!

    try:
        with pdfplumber.open(input_path) as pdf:
            current_row = 1
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        if table:
                            for row in table:
                                for col_idx, cell_value in enumerate(row, 1):
                                    # نكتب النص كما هو دون أي تغيير
                                    ws.cell(row=current_row, column=col_idx, value=cell_value)
                                current_row += 1
                            current_row += 1
                else:
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            ws.cell(row=current_row, column=1, value=line)
                            current_row += 1
    except Exception as e:
        logger.error(f"PDF to Excel extraction error: {e}")
        raise

    # تنسيق الخلايا
    header_font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    cell_font = Font(name='Arial', size=12)
    cell_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 20)):
        for cell in row:
            cell.font = cell_font
            cell.alignment = cell_alignment
            cell.border = thin_border

    wb.save(output_path)
   
    
@router.callback_query()
async def handle_conversion_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    data = callback.data
    
    conversion_map = {
        "convert_word2pdf": ("docx", "pdf", "Word → PDF"),
        "convert_pdf2word": ("pdf", "docx", "PDF → Word"),
        "convert_excel2pdf": ("xlsx", "pdf", "Excel → PDF"),
        "convert_pdf2excel": ("pdf", "xlsx", "PDF → Excel"),
        "convert_excel2word": ("xlsx", "docx", "Excel → Word"),
        "convert_word2excel": ("docx", "xlsx", "Word → Excel"),
    }
    
    if data == "convert_any":
        user_conversion_choice[user_id] = ("any", None, "أي صيغة لأي صيغة")
        await callback.message.answer(
            "📁 *أرسل الملف الذي تريد تحويله*",
            parse_mode="Markdown"
        )
        await callback.answer("تم")
        return
    
    if data in conversion_map:
        source, target, label = conversion_map[data]
        user_conversion_choice[user_id] = (source, target, label)
        
        await callback.message.answer(
            f"📁 *{label}*\n\n"
            f"أرسل ملف *{source.upper()}* ليتم تحويله إلى *{target.upper()}*",
            parse_mode="Markdown"
        )
        await callback.answer("تم")

@router.message(Command("start"))
async def cmd_start(message: types.Message):
    update_user_activity(message.from_user)
    
    keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="💬 ابدأ محادثة"), KeyboardButton(text="🖼️ تحليل صورة")],
            [KeyboardButton(text="📄 تحويل نص لملف"), KeyboardButton(text="📊 تحويل لإكسيل")],
            [KeyboardButton(text="🎤 إرسال صوت"), KeyboardButton(text="🌐 ترجمة فورية")],
            [KeyboardButton(text="🔄 تحويل ملفات"), KeyboardButton(text="👨‍💻 تواصل مع المبرمج")]
        ],
        resize_keyboard=True,
        input_field_placeholder="اختر من القائمة..."
    )
    
    await message.answer(
        "🎉 أهلاً بك! أنا مستشار الذكاء الاصطناعي الخارق.\n\n"
        "✨ ماذا يمكنني أن أفعل لك؟\n"
        "- الإجابة عن أي سؤال\n"
        "- كتابة وشرح الأكواد البرمجية\n"
        "- تحويل النصوص إلى Word أو PDF أو Excel\n"
        "- تحويل الملفات بين الصيغ\n"
        "- تحليل الصور والمستندات\n"
        "- الاستماع إلى الرسائل الصوتية\n"
        "- الترجمة الفورية لأي لغة\n"
        "- تصميم برومبت احترافي للصور\n\n"
        "💬 تحدث معي طبيعياً وسأفهمك!\n\n"
        "━━━━━━━━━━━━━━━━━━\n"
        f"👨‍💻 المبرمج: {DEVELOPER_NAME}\n"
        "━━━━━━━━━━━━━━━━━━",
        reply_markup=keyboard
    )

@router.message(Command("admin"))
async def cmd_admin(message: types.Message):
    update_user_activity(message.from_user)
    
    if message.from_user.id != ADMIN_USER_ID:
        await message.answer("⛔ عذراً، هذا الأمر متاح فقط لمالك البوت.")
        return
    
    conn = sqlite3.connect('bot_stats.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT COUNT(*) FROM users')
    total_users = cursor.fetchone()[0]
    
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute('SELECT active_users, total_messages FROM daily_stats WHERE date=?', (today,))
    today_stats = cursor.fetchone() or (0, 0)
    
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    cursor.execute('SELECT active_users, total_messages FROM daily_stats WHERE date=?', (yesterday,))
    yesterday_stats = cursor.fetchone() or (0, 0)
    
    cursor.execute('SELECT SUM(total_messages) FROM daily_stats')
    total_messages_all_time = cursor.fetchone()[0] or 0
    
    one_day_ago = (datetime.now() - timedelta(hours=24)).strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute('SELECT COUNT(*) FROM users WHERE last_active >= ?', (one_day_ago,))
    online_users = cursor.fetchone()[0]
    
    offline_users = total_users - online_users
    
    cursor.execute('SELECT username, first_name, last_active FROM users ORDER BY last_active DESC LIMIT 5')
    recent_users = cursor.fetchall()
    
    conn.close()
    
    stats_message = (
        "📊 *لوحة الإحصائيات*\n\n"
        f"👥 إجمالي المستخدمين: {total_users}\n"
        f"🟢 متصل: {online_users}\n"
        f"🔴 غير متصل: {offline_users}\n\n"
        f"📅 اليوم: {today_stats[0]} نشط | {today_stats[1]} رسالة\n"
        f"📆 أمس: {yesterday_stats[0]} نشط | {yesterday_stats[1]} رسالة\n\n"
        f"💬 إجمالي الرسائل: {total_messages_all_time}"
    )
    
    await message.answer(stats_message, parse_mode="Markdown")

@router.message(Command("reset"))
async def cmd_reset(message: types.Message):
    update_user_activity(message.from_user)
    gemini_client.conversations.pop(str(message.from_user.id), None)
    await message.answer("🔄 تم مسح سياق المحادثة.")

@router.message(Command("translate"))
async def cmd_translate(message: types.Message):
    update_user_activity(message.from_user)
    await message.answer(
        "🌐 *الترجمة الفورية*\n\n"
        "يمكنك الترجمة بطريقتين:\n\n"
        "1️⃣ *أرسل النص بهذا الشكل:*\n"
        "`ترجم إلى الفرنسية: مرحباً، كيف حالك؟`\n\n"
        "2️⃣ *أرسل رسالة صوتية:*\n"
        "سأحولها إلى نص ثم أترجمها لك.\n\n"
        "📝 *مثال للأوامر:*\n"
        "- ترجم إلى الإنجليزية: النص\n"
        "- ترجم إلى الإسبانية: النص\n"
        "- ترجم إلى الألمانية: النص",
        parse_mode="Markdown"
    )

@router.message(F.text.in_({"💬 ابدأ محادثة", "🖼️ تحليل صورة", "📄 تحويل نص لملف", "📊 تحويل لإكسيل", "🎤 إرسال صوت", "👨‍💻 تواصل مع المبرمج", "🔄 تحويل ملفات", "🌐 ترجمة فورية"}))
async def handle_buttons(message: types.Message):
    update_user_activity(message.from_user)
    
    if message.text == "💬 ابدأ محادثة":
        await message.answer("📝 أنا جاهز! أرسل سؤالك أو طلبك وسأجيبك فوراً.")
    elif message.text == "🖼️ تحليل صورة":
        await message.answer("🖼️ أرسل لي الصورة التي تريد تحليلها.")
    elif message.text == "📄 تحويل نص لملف":
        await message.answer(
            "📄 أرسل لي النص الذي تريد تحويله.\n\n"
            "• *وورد:* حولي النص دا لملف وورد: ...\n"
            "• *PDF:* حولي النص دا لملف PDF: ...\n"
            "• *اكسيل:* حولي النص دا لملف اكسيل: ...",
            parse_mode="Markdown"
        )
    elif message.text == "📊 تحويل لإكسيل":
        await message.answer(
            "📊 أرسل لي النص الذي تريد تحويله إلى ملف Excel.\n\n"
            "مثال: *حولي النص دا لملف اكسيل: الاسم, العمر, المدينة\nأحمد, 25, القاهرة*",
            parse_mode="Markdown"
        )
    elif message.text == "🎤 إرسال صوت":
        await message.answer(
            "🎤 أرسل لي رسالة صوتية وسأقوم بما يلي:\n\n"
            "1️⃣ تحويلها إلى نص مكتوب\n"
            "2️⃣ الرد على محتواها\n"
            "3️⃣ يمكنك أيضاً طلب إنشاء ملف Word أو PDF أو Excel من النص المستخرج"
        )
    elif message.text == "👨‍💻 تواصل مع المبرمج":
        await message.answer(
            f"👨‍💻 *المبرمج:* {DEVELOPER_NAME}\n\n"
            f"📧 *للتواصل:* \u200E@omarawad68",
            parse_mode="Markdown"
        )
    elif message.text == "🔄 تحويل ملفات":
        await message.answer(
            "🔄 *اختر نوع التحويل الذي تريده:*",
            parse_mode="Markdown",
            reply_markup=get_conversion_keyboard()
        )
    elif message.text == "🌐 ترجمة فورية":
        await message.answer(
            "🌐 *الترجمة الفورية*\n\n"
            "يمكنك الترجمة بطريقتين:\n\n"
            "1️⃣ *أرسل النص بهذا الشكل:*\n"
            "`ترجم إلى الفرنسية: مرحباً، كيف حالك؟`\n\n"
            "2️⃣ *أرسل رسالة صوتية:*\n"
            "سأحولها إلى نص ثم أترجمها لك.\n\n"
            "📝 *مثال للأوامر:*\n"
            "- ترجم إلى الإنجليزية: النص\n"
            "- ترجم إلى الإسبانية: النص\n"
            "- ترجم إلى الألمانية: النص",
            parse_mode="Markdown"
        )

@router.message(F.text)
async def handle_message(message: types.Message):
    update_user_activity(message.from_user)
    user_text = message.text
    text_lower = user_text.lower()

    if user_text in ["💬 ابدأ محادثة", "🖼️ تحليل صورة", "📄 تحويل نص لملف", "📊 تحويل لإكسيل", "🎤 إرسال صوت", "👨‍💻 تواصل مع المبرمج", "🔄 تحويل ملفات", "🌐 ترجمة فورية"]:
        return

    # ==================== التحقق من الملفات المعلقة ====================
    user_id = message.from_user.id
    if user_id in user_pending_file:
        chosen_format = None
        if user_text.lower() in ['pdf']: chosen_format = 'pdf'
        elif user_text.lower() in ['word', 'docx']: chosen_format = 'docx'
        elif user_text.lower() in ['excel', 'xlsx']: chosen_format = 'xlsx'
        
        if chosen_format:
            pending = user_pending_file.pop(user_id)
            await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
            try:
                inpath = f"/tmp/{user_id}_{pending['filename']}"
                with open(inpath, 'wb') as f:
                    f.write(pending['file_bytes'])
                expected_out = f"/tmp/{os.path.splitext(pending['filename'])[0]}.{chosen_format}"
                
                if chosen_format == 'xlsx' and inpath.lower().endswith('.pdf'):
                    convert_pdf_to_excel(inpath, expected_out)
                else:
                    run_libreoffice(['--convert-to', chosen_format, '--outdir', '/tmp/', inpath])
                
                if os.path.exists(expected_out) and os.path.getsize(expected_out) > 100:
                    await message.reply_document(
                        FSInputFile(expected_out),
                        caption=f"✅ تم التحويل إلى {chosen_format.upper()}"
                    )
                else:
                    possible_files = glob.glob(f"/tmp/*.{chosen_format}")
                    found = False
                    for pf in possible_files:
                        if os.path.getsize(pf) > 100:
                            await message.reply_document(
                                FSInputFile(pf),
                                caption=f"✅ تم التحويل إلى {chosen_format.upper()}"
                            )
                            os.remove(pf)
                            found = True
                            break
                    if not found:
                        await message.reply("❌ فشل التحويل.")
                if os.path.exists(inpath): os.remove(inpath)
                if os.path.exists(expected_out): os.remove(expected_out)
            except Exception as e:
                logger.error(f"Convert error: {e}")
                await message.reply("❌ حدث خطأ أثناء التحويل.")
            return

    intent, content = detect_conversion_intent(user_text)
    
    if intent == "EXCEL_NEED_TEXT":
        return await message.reply("📊 ما هو النص الذي تريد تحويله إلى ملف Excel؟")
    if intent == "WORD_NEED_TEXT":
        return await message.reply("📝 ما هو النص الذي تريد تحويله إلى ملف Word؟")
    if intent == "PDF_NEED_TEXT":
        return await message.reply("📕 ما هو النص الذي تريد تحويله إلى ملف PDF؟")
    
    if intent == "excel" and content:
        await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
        try:
            path = f"/tmp/{message.from_user.id}_doc.xlsx"
            create_excel_file(content, path)
            await message.reply_document(FSInputFile(path), caption="📊 ملف Excel جاهز!")
            os.remove(path)
            return
        except Exception as e:
            logger.error(f"Excel error: {e}")
            return await message.reply("❌ حدث خطأ في إنشاء ملف Excel.")
    
    if intent == "docx" and content:
        await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
        try:
            path = f"/tmp/{message.from_user.id}_doc.docx"
            create_docx_file(content, path)
            await message.reply_document(FSInputFile(path), caption="📄 ملف Word جاهز!")
            os.remove(path)
            return
        except Exception as e:
            logger.error(f"Word error: {e}")
            return await message.reply("❌ حدث خطأ في إنشاء ملف Word.")
            
    if intent == "pdf" and content:
        await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
        try:
            path = f"/tmp/{message.from_user.id}_doc.pdf"
            create_pdf_file(content, path)
            await message.reply_document(FSInputFile(path), caption="📕 ملف PDF جاهز!")
            os.remove(path)
            return
        except Exception as e:
            logger.error(f"PDF error: {e}")
            return await message.reply("❌ حدث خطأ في إنشاء ملف PDF.")

    image_keywords = ["اعملي صورة", "اعمل صورة", "ارسم", "صمملي", "تخيل", "صورلي", "توليد صورة", "انشاء صورة", "صمم صورة"]
    is_image_request = any(keyword in text_lower for keyword in image_keywords)

    if is_image_request:
        await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
        prompt_request = f"""حوّل الطلب التالي إلى أمر (Prompt) إبداعي واحترافي باللغة العربية لاستخدامه مع مولدات الصور بالذكاء الاصطناعي. أضف تفاصيل عن الإضاءة، الألوان، الزاوية، والجو العام.
        
        طلب المستخدم: {user_text}
        
        اكتب فقط نص الأمر (البرومبت) بدون أي مقدمات أو شرح إضافي."""
        
        generated_prompt = await gemini_client.generate(prompt_request, str(message.from_user.id))
        
        final_response = f"🎨 *تم تصميم برومبت احترافي لطلبك:*\n\n`{generated_prompt}`\n\n🖼️ يمكنك نسخ هذا النص ولصقه في أي أداة لتوليد الصور بالذكاء الاصطناعي."
        await message.reply(final_response, parse_mode="Markdown")
        return

    # ==================== الترجمة الفورية ====================
    translate_triggers = ["ترجم إلى", "ترجم الى", "ترجم لـ", "ترجمة إلى", "ترجمة لـ", "translate to"]
    is_translate_request = any(trigger in text_lower for trigger in translate_triggers)

    if is_translate_request:
        target_lang = None
        text_to_translate = None
        
        for trigger in translate_triggers:
            if trigger in text_lower:
                idx = text_lower.find(trigger)
                rest = user_text[idx + len(trigger):].strip()
                
                if ':' in rest:
                    target_lang, text_to_translate = rest.split(':', 1)
                    target_lang = target_lang.strip()
                    text_to_translate = text_to_translate.strip()
                else:
                    target_lang = rest.strip()
                break
        
        if target_lang and text_to_translate:
            await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
            
            prompt = f"ترجم النص التالي إلى {target_lang}. أرسل الترجمة فقط بدون أي كلام إضافي:\n\n{text_to_translate}"
            translation = await gemini_client.generate(prompt, str(message.from_user.id))
            
            await message.answer(f"🌐 *الترجمة إلى {target_lang}:*\n\n{translation}", parse_mode="Markdown")
            return
        elif target_lang:
            await message.reply(f"🌐 *من فضلك أرسل النص الذي تريد ترجمته إلى {target_lang}.*\n\nمثال: *ترجم إلى {target_lang}: النص هنا*", parse_mode="Markdown")
            return

    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
    resp = await gemini_client.generate(user_text, str(message.from_user.id))
    for i in range(0, len(resp), 4000):
        await message.answer(resp[i:i+4000])

@router.message(F.photo)
async def handle_photo(message: types.Message, bot: Bot):
    update_user_activity(message.from_user)
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    try:
        photo = message.photo[-1]
        file_info = await bot.get_file(photo.file_id)
        bio = BytesIO()
        await bot.download_file(file_info.file_path, bio)
        bio.seek(0)
        img_bytes, mime = convert_image_to_png(bio.read())
        b64 = base64.b64encode(img_bytes).decode()
        caption = message.caption or "حلل هذه الصورة"
        resp = await gemini_client.generate_with_media(caption, [
            {"inline_data": {"mime_type": mime, "data": b64}}
        ])
        for i in range(0, len(resp), 4000):
            await message.reply(resp[i:i+4000])
    except Exception as e:
        logger.error(f"Photo error: {e}")
        await message.reply("عذراً، حدث خطأ.")

@router.message(F.document)
async def handle_document(message: types.Message, bot: Bot):
    update_user_activity(message.from_user)
    doc = message.document
    fname = doc.file_name or "مستند.xlsx"
    mime = doc.mime_type or ""
    cap = message.caption or ""
    user_id = message.from_user.id
    
    # التحقق مما إذا كان المستخدم قد اختار "any" (أي صيغة لأي صيغة)
    if user_id in user_conversion_choice and user_conversion_choice[user_id][0] == "any":
        target = None
        if cap:
            c = cap.lower()
            if "pdf" in c: target = "pdf"
            elif "word" in c or "docx" in c: target = "docx"
            elif "excel" in c or "xlsx" in c: target = "xlsx"
        
        if target:
            await bot.send_chat_action(chat_id=message.chat.id, action="typing")
            try:
                file_info = await bot.get_file(doc.file_id)
                file_bytes = await bot.download_file(file_info.file_path)
                inpath = f"/tmp/{user_id}_{fname}"
                with open(inpath, 'wb') as f:
                    f.write(file_bytes.read())
                expected_out = f"/tmp/{os.path.splitext(fname)[0]}.{target}"
                
                if target == 'xlsx' and inpath.lower().endswith('.pdf'):
                    convert_pdf_to_excel(inpath, expected_out)
                else:
                    run_libreoffice(['--convert-to', target, '--outdir', '/tmp/', inpath])
                
                if os.path.exists(expected_out) and os.path.getsize(expected_out) > 100:
                    await message.reply_document(
                        FSInputFile(expected_out),
                        caption=f"✅ تم التحويل إلى {target.upper()}"
                    )
                else:
                    possible_files = glob.glob(f"/tmp/*.{target}")
                    found = False
                    for pf in possible_files:
                        if os.path.getsize(pf) > 100:
                            await message.reply_document(
                                FSInputFile(pf),
                                caption=f"✅ تم التحويل إلى {target.upper()}"
                            )
                            os.remove(pf)
                            found = True
                            break
                    if not found:
                        await message.reply("❌ فشل التحويل.")
                if os.path.exists(inpath): os.remove(inpath)
                if os.path.exists(expected_out): os.remove(expected_out)
            except Exception as e:
                logger.error(f"Convert error: {e}")
                await message.reply("❌ حدث خطأ أثناء التحويل.")
            if user_id in user_conversion_choice:
                del user_conversion_choice[user_id]
            return
        else:
            # تخزين الملف مؤقتاً وانتظار اختيار الصيغة
            file_info = await bot.get_file(doc.file_id)
            file_bytes = await bot.download_file(file_info.file_path)
            user_pending_file[user_id] = {
                'file_bytes': file_bytes.read(),
                'filename': fname
            }
            await message.reply("📝 *إلى أي صيغة تريد التحويل؟*\n• pdf\n• word\n• excel", parse_mode="Markdown")
            return
    
    target = None
    if user_id in user_conversion_choice:
        source, target, label = user_conversion_choice[user_id]
    
    if not target and cap:
        c = cap.lower()
        if "pdf" in c: target = "pdf"
        elif "word" in c or "docx" in c: target = "docx"
        elif "excel" in c or "xlsx" in c: target = "xlsx"
        elif "ppt" in c or "pptx" in c: target = "pptx"
    
    if target:
        await bot.send_chat_action(chat_id=message.chat.id, action="typing")
        try:
            file_info = await bot.get_file(doc.file_id)
            file_bytes = await bot.download_file(file_info.file_path)
            inpath = f"/tmp/{user_id}_{fname}"
            with open(inpath, 'wb') as f:
                f.write(file_bytes.read())
            expected_out = f"/tmp/{os.path.splitext(fname)[0]}.{target}"
            
            if target == 'xlsx' and inpath.lower().endswith('.pdf'):
                convert_pdf_to_excel(inpath, expected_out)
            else:
                run_libreoffice(['--convert-to', target, '--outdir', '/tmp/', inpath])
            
            if os.path.exists(expected_out) and os.path.getsize(expected_out) > 100:
                await message.reply_document(
                    FSInputFile(expected_out),
                    caption=f"✅ تم التحويل إلى {target.upper()}"
                )
            else:
                possible_files = glob.glob(f"/tmp/*.{target}")
                found = False
                for pf in possible_files:
                    if os.path.getsize(pf) > 100:
                        await message.reply_document(
                            FSInputFile(pf),
                            caption=f"✅ تم التحويل إلى {target.upper()}"
                        )
                        os.remove(pf)
                        found = True
                        break
                if not found:
                    await message.reply("❌ فشل التحويل.")
            if os.path.exists(inpath): os.remove(inpath)
            if os.path.exists(expected_out): os.remove(expected_out)
        except Exception as e:
            logger.error(f"Convert error: {e}")
            await message.reply("❌ حدث خطأ أثناء التحويل.")
        return
    
    supported = [
        "application/pdf", "text/plain",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "text/csv"
    ]
    if mime not in supported:
        return await message.reply("⚠️ نوع الملف غير مدعوم.")
    
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    try:
        info = await bot.get_file(doc.file_id)
        bio = BytesIO()
        await bot.download_file(info.file_path, bio)
        bio.seek(0)
        fb = bio.read()
        text = ""
        if mime == "text/plain" or mime == "text/csv":
            text = fb.decode('utf-8', errors='ignore')
        elif mime == "application/pdf":
            import PyPDF2
            r = PyPDF2.PdfReader(BytesIO(fb))
            for p in r.pages:
                text += p.extract_text() or ""
        elif "word" in mime:
            import docx as dx
            dxf = dx.Document(BytesIO(fb))
            text = "\n".join([p.text for p in dxf.paragraphs])
        elif "excel" in mime or "spreadsheet" in mime:
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(fb), read_only=True)
            ws = wb.active
            for row in ws.iter_rows(values_only=True):
                text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
        if not text.strip():
            return await message.reply("⚠️ لم أستطع استخراج نص.")
        prompt = f"حلل هذا المستند ({fname}). {cap or 'قدم ملخصاً'}\n\n{text[:10000]}"
        resp = await gemini_client.generate(prompt, str(user_id))
        for i in range(0, len(resp), 4000):
            await message.reply(resp[i:i+4000])
    except Exception as e:
        logger.error(f"Doc error: {e}")
        await message.reply("عذراً، حدث خطأ.")

@router.message(F.voice)
async def handle_voice(message: types.Message, bot: Bot):
    update_user_activity(message.from_user)
    await bot.send_chat_action(chat_id=message.chat.id, action="typing")
    
    ogg_path = f"/tmp/{message.from_user.id}_voice.ogg"
    wav_path = f"/tmp/{message.from_user.id}_voice.wav"
    
    try:
        voice = message.voice
        file_info = await bot.get_file(voice.file_id)
        bio = BytesIO()
        await bot.download_file(file_info.file_path, bio)
        bio.seek(0)
        
        with open(ogg_path, "wb") as f:
            f.write(bio.read())
        bio.close()
        
        try:
            subprocess.run(
                ['ffmpeg', '-i', ogg_path, '-ar', '16000', '-ac', '1', wav_path],
                check=True, capture_output=True, timeout=30
            )
            logger.info("Audio converted to WAV")
        except Exception as e:
            logger.error(f"ffmpeg error: {e}")
            await message.reply("🎤 عذراً، فشل تحويل الصوت.")
            return
        
        import speech_recognition as sr
        recognizer = sr.Recognizer()
        
        with sr.AudioFile(wav_path) as source:
            audio = recognizer.record(source)
        
        text = None
        for lang in ["ar-AR", "en-US", ""]:
            try:
                text = recognizer.recognize_google(audio, language=lang) if lang else recognizer.recognize_google(audio)
                if text: break
            except sr.UnknownValueError:
                continue
            except sr.RequestError as e:
                logger.error(f"Google API error: {e}")
                await message.reply("⚠️ خدمة التعرف على الصوت غير متاحة حالياً.")
                return
        
        if not text:
            await message.reply("🎤 لم أتمكن من فهم الصوت.")
            return
        
        await message.reply(f"🎤 *لقد فهمت:* _{text}_", parse_mode="Markdown")
        resp = await gemini_client.generate(text, str(message.from_user.id))
        for i in range(0, len(resp), 4000):
            await message.answer(resp[i:i+4000])
            
    except ImportError:
        await message.reply("⚠️ مكتبة الصوت غير مثبتة.")
    except Exception as e:
        logger.error(f"Voice error: {e}")
        await message.reply("🎤 عذراً، حدث خطأ.")
    finally:
        if os.path.exists(ogg_path): os.remove(ogg_path)
        if os.path.exists(wav_path): os.remove(wav_path)

async def handle_web_chat(request):
    try:
        data = await request.json()
        user_text = data.get('content', '')
        user_id = request.headers.get('X-User-Id', 'web_user')
        
        if not user_text:
            return web.json_response({'status': 'error', 'message': 'نص فارغ'})
        
        response = await gemini_client.generate(user_text, user_id)
        
        return web.json_response({
            'status': 'success',
            'response': response
        })
    except Exception as e:
        logger.error(f"Web chat error: {e}")
        return web.json_response({'status': 'error', 'message': 'حدث خطأ'})

async def init_web_server():
    app = web.Application()
    app.router.add_post('/api/chat', handle_web_chat)
    
    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, '0.0.0.0', 8000)
    await site.start()
    logger.info("Web server started on port 8000")

async def main():
    init_db()
    bot = Bot(token=os.getenv("TELEGRAM_BOT_TOKEN"))
    dp = Dispatcher()
    dp.include_router(router)
    
    await init_web_server()
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
